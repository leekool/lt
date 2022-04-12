from datetime import date
import subprocess
import shutil
import json
import click
import docx
import os
import sys
import os.path
import time
import re
import pywinauto
from openpyxl import load_workbook

dt = date.today()
rasdial = subprocess.check_output('rasdial').decode('utf-8')  # for checking VPN connection
CONTEXT_SETTINGS = dict(help_option_names=['-h', '--help', '--h'])

# reads config.json
with open('config.json', 'r') as jsonfile:
    config = json.load(jsonfile)
    jsonfile.close()


@click.group(context_settings=CONTEXT_SETTINGS)
def cli():
    pass


# prints the running sheet last read from the 'daily' command
@cli.command(name='sheet', help='Prints the last read running sheet.')
def sheet():
    click.echo(f'\n{config["sheet"]}')


# change prefix manually (document name before turn number)
@cli.command(name='prefix', help='Change the value of \'prefix\' (the beginning of a turn name).')
@click.argument('name')
def prefix(name):
    previous = config['prefix']
    config['prefix'] = name
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)
    click.echo(f'\nChanged \'{previous}\' to \'{name}\'.')


# change intials (used in 'daily' when reading running sheet)
@cli.command(name='initials', help='Change the value of \'initials\' (used in \'daily\' when reading running sheet).')
@click.argument('name')
def initials(name):
    previous = config['initials']
    config['initials'] = name.upper()
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)
    click.echo(f'\nChanged \'{previous}\' to \'{name.upper()}\'.')


# change speaker names
@cli.command(name='s1', help='Change the value of \'speaker1\'.')
@click.argument('name')
def s1(name):
    previous = config['speaker1']
    config['speaker1'] = name
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)
    click.echo(f'\nChanged \'{previous}\' to \'{name}\'.')


@cli.command(name='s2', help='Change the value of \'speaker2\'.')
@click.argument('name')
def s2(name):
    previous = config['speaker2']
    config['speaker2'] = name
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)
    click.echo(f'\nChanged \'{previous}\' to \'{name}\'.')


@cli.command(name='s3', help='Change the value of \'speaker3\'.')
@click.argument('name')
def s3(name):
    previous = config['speaker3']
    config['speaker3'] = name
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)
    click.echo(f'\nChanged \'{previous}\' to \'{name}\'.')


@cli.command(name='s4', help='Change the value of \'speaker4\'.')
@click.argument('name')
def s4(name):
    previous = config['speaker4']
    config['speaker4'] = name
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)
    click.echo(f'\nChanged \'{previous}\' to \'{name}\'.')


# create and open new doc
@cli.command(name='doc',
             help='Creates and opens a Word document corresponding with turn specified.')
@click.option('--r',
              is_flag=True,
              help='Adds notation for turns that start on a resumption.')
@click.argument('turn')
def doc(turn, r):
    config['last_turn'] = config['prefix'] + turn.upper()
    doc = docx.Document('C:/Users/LEE/AppData/Roaming/Microsoft/Templates/AGNSW 2021.docx')
    doc._body.clear_content()

    options = {'a': '[10.00 - 10.15]',
               'b': '[10.15 - 10.30]',
               'c': '[10.30 - 10.45]',
               'd': '[10.45 - 11.00]',
               'e': '[11.00 - 11.15]',
               'f': '[11.15 - 11.30]',
               'g': '[11.30 - 11.45]',
               'h': '[11.45 - 12.00]',
               'i': '[12.00 - 12.15]',
               'j': '[12.15 - 12.30]',
               'k': '[12.30 - 12.45]',
               'l': '[12.45 - 1.00]',
               'l2': '[1.00 - 1.15]',
               'l3': '[1.15 - 1.30]',
               'l4': '[1.30 - 1.45]',
               'm': '[2.00 - 2.15]',
               'n': '[2.15 - 2.30]',
               'o': '[2.30 - 2.45]',
               'p': '[2.45 - 3.00]',
               'q': '[3.00 - 3.15]',
               'r': '[3.15 - 3.30]',
               's': '[3.30 - 3.45]',
               't': '[3.45 - 4.00]',
               'u': '[4.00 - 4.15]', }

    if turn in options:
        doc.add_paragraph(f'START OF TURN {config["last_turn"]} {options[turn]}')
    else:
        sys.exit('\nInvalid turn.')

    # save turn's path to config.json so that it can be copied to VPN drive when finished
    config['last_turn_path'] = f'{config["working_path"]}{config["last_turn"]}.docx'
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)
    doc.save(config['last_turn_path'])
    click.echo(f'\nCreated Word document: {config["last_turn"]}.docx')

    # open document
    os.startfile(f'{config["working_path"]}{config["last_turn"]}.docx')
    try:
        app = pywinauto.Application().connect(best_match=config['last_turn'], timeout=5).top_window()
    except pywinauto.timings.TimeoutError:
        sys.exit(f'\n{config["last_turn"]}.docx did not open.')
    else:
        app.type_keys('{END}')  # moves cursor to end of document

    # if --r option is used (resumption)
    if r:
        app.type_keys('{SPACE}[RESUMPTION]')


# open daily folders and write path to config.json
@cli.command(name='daily',
             help='Open folders relevant to today\'s date and presiding officer specified. Writes \'daily_path\' to config.json.')
def daily():
    if 'Legal Transcripts' in rasdial:  # checks if connected to VPN
        pass
    else:
        sys.exit('\nNot connected to \'Legal Transcripts VPN 2\'.')

    # creates list of folders in folders for today's date
    list = []
    for parent, dirs, files in os.walk(f'X:/{dt.strftime("%Y")}/{dt.strftime("%B")}/{dt.strftime("%d.%m.%y")}/'):
        for dirname in dirs:
            list.append(dirname)  # creates list of folders in path

    click.echo()  # blank line - probably a better way to do this

    # numbers folders in 'list' and allows you to pick a folder by inputting a number
    for cnt, name in enumerate(list, 1):
        sys.stdout.write('%d. %s\n\r' % (cnt, name))
    choice = int(input('\nSelect daily folder [1-%s]: ' % cnt)) - 1

    # checks 'daily_path' exists
    config['daily_path'] = f'X:/{dt.strftime("%Y")}/{dt.strftime("%B")}/{dt.strftime("%d.%m.%y")}/{list[choice]}/'
    if os.path.exists(config['daily_path']):
        click.echo(f'\nChanged daily path to \'{config["daily_path"]}\'.')
    else:
        sys.exit('\nFolder doesn\'t exist.')

    # attempts to find running sheet in 'daily_path'
    rs = [s for s in os.listdir(config['daily_path']) if 'running' in s or list[choice] in s]
    if rs == []:  # if no match
        os.startfile(config['daily_path'])
        sys.exit('\nRunning sheet not found.  Enter \'prefix\' manually.')
    else:
        shutil.copy(f'{config["daily_path"]}{rs[0]}', f'C:/Users/LEE/Desktop/{rs[0]}')
        
    # if running sheet is .doc attempts to convert it to .docx
    while rs[0].endswith('.doc'):
        os.startfile(f'{config["working_path"]}{rs[0]}')
        app = pywinauto.Application().connect(best_match=rs[0], timeout=5)
        app.top_window().type_keys('^+s')  # opens save as... dialog
        dlg = app.window(class_name='#32770')  # connects to save as... dialog
        dlg.ComboBox2.select('Word Document ')  # selects .docx in dropdown
        dlg.Button8.click()  # clicks save
        app = pywinauto.Application().connect(best_match=rs[0], timeout=5).top_window()  # connect to new .docx
        app.close()
        os.remove(f'{config["working_path"]}{rs[0]}')  # deletes .doc running sheet from desktop
        rs[0] = re.sub('.doc', '.docx', rs[0])  # need to read regex docs and do this properly

    doc = docx.Document(f'{config["working_path"]}{rs[0]}')
    table = doc.tables[0]
    data = []

    # gets text from table rows and puts it into list (data)
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        data.append(' '.join(text))

    # finds 'prefix' in str(data)
    previous = config['prefix']
    config['prefix'] = re.search(rf'\b{dt.strftime("%d%m")}\w+', str(data))  # finds word containg today's date
    config['prefix'] = re.sub(r'[A-Z]', '', config['prefix'].group())  # removes capital letters (turn letter)
    click.echo(f'\nChanged prefix from \'{previous}\' to \'{config["prefix"]}\'.')

    # finds rows containing 'intials' and prints
    turns = [i for i in data if config['initials'] in i]
    config['sheet'] = '\n'.join(turns)
    click.echo(f'\n{config["sheet"]}')  # prints turns corresponding with initials

    # saves 'daily_path', 'prefix', and 'sheet' to config.json
    with open('config.json', 'w') as jsonfile:
        json.dump(config, jsonfile)

    # open sound folder and delete running sheet
    os.startfile(f'S:/AGNSW DAILIES/{dt.strftime("%Y%m%d")}')
    os.remove(f'{config["working_path"]}{rs[0]}')  # deletes .docx running sheet from desktop


# connect/disconnect VPN
@cli.command(name='vpn',
             help='Toggles connection to \'Legal Transcipts VPN 2\'.')
def vpn():
    if 'No connections' in rasdial:
        click.echo('\nNot connected.\n\nConnecting...')
        os.system('start rasphone')
        app = pywinauto.Application().connect(title='Network Connections', timeout=3).top_window()
        app.type_keys('{ENTER 2}')
    else:
        click.echo('\nConnected.\n\nDisconnecting...')
        os.system('rasphone -h "Legal Transcripts VPN 2"')


# saves and copies 'last_turn' to 'daily_path' and saves info in excel
@cli.command(name='save',
             help='Saves and closes \'last_turn\', writes info to Excel invoice, and moves it to \'daily_path\'.')
def save():
    try:
        app = pywinauto.Application().connect(best_match=config['last_turn'], timeout=2).top_window()
    except pywinauto.timings.TimeoutError:
        click.echo(f'\n{config["last_turn"]}.docx is not open.')
    else:
        app.type_keys('^s')  # saves the document
        app.close()
        time.sleep(0.2)
        doc = docx.Document(config['last_turn_path'])
        doc.add_paragraph()  # adds blank line - probably a better way to do this
        doc.add_paragraph(f'END OF TURN {config["last_turn"]}')
        doc.save(config['last_turn_path'])
        click.echo(f'\nSaved and closed: {config["last_turn"]}.docx')

    #  counts words in document
    word_count = 0
    for para in doc.paragraphs:
        if para.text.find('--') >= 0:  # accounts for microsoft word counting breaks as words
            word_count += 1
        word_count = word_count + len(para.text.split())
    click.echo(f'\nCounted {word_count} words.')

    wb = load_workbook(filename='C:/Users/LEE/Documents/work/Lee Luppi transcription invoice period end 15.04.22.xlsx')

    #  finds next empty row in excel invoice between rows 15-81
    empty_row = 0
    for row in wb.active.iter_rows(min_row=15, max_row=81, max_col=1):
        for cell in row:
            if cell.value is None:
                empty_row = cell.row
        if empty_row >= 1:
            break

    #  writes 'last_turn', date, and 'word_count' to their columns
    wb.active.cell(row=empty_row, column=1).value = config['last_turn']
    wb.active.cell(row=empty_row, column=2).value = dt.strftime('%d.%m.%y')
    wb.active.cell(row=empty_row, column=4).value = word_count
    wb.save('C:/Users/LEE/Documents/work/Lee Luppi transcription invoice period end 15.04.22.xlsx')
    click.echo(f'\nCopied \'{config["last_turn"]}\', \'{dt.strftime("%d.%m.%y")}\', and \'{word_count}\' to row {empty_row}.')

    #  moves document to 'daily_path'
    if 'Legal Transcripts' in rasdial:  # checks if connected to VPN
        shutil.move(config['last_turn_path'], config['daily_path'])
        click.echo(f'\n{config["last_turn"]}.docx moved to \'{config["daily_path"]}\'.')
    else:
        click.echo('\nNot connected to VPN.  Could not move document to daily folder.')


if __name__ == '__main__':
    cli()
